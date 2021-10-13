# -*- coding: utf-8 -*-
from docx import Document

from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.oxml.ns import qn # 中文格式

from docx.shared import Pt # 磅数

from docx.shared import Inches # 图片尺寸

from bs4 import BeautifulSoup

from urllib.request import urlopen

import pandas as pd

import numpy as np

import matplotlib.pyplot as plt

import os

import matplotlib as mpl

mpl.rcParams['font.sans-serif'] = ['KaiTi', 'SimHei', 'FangSong']  #汉字字体,优先使用楷体，如果找不到楷体，则使用黑体
mpl.rcParams['font.size'] = 12  # 字体大小
mpl.rcParams['axes.unicode_minus'] = False





c_name = input('请输入公司名称：')

stock_code = input("请输入股票代码：")

st = input('请输入开始年份：')

et = input('请输入结束年份：')



data_list =['zcfzb','lrb','xjllb']

adrees ='http://quotes.money.163.com'

y_list =['报告日期']+[str(m)+'年' for m in list(range(int(st)-1,int(et)+1))][::-1]





for i in data_list:

    url ='http://quotes.money.163.com/f10/'+i+'_'+stock_code+'.html?type=year'



    html = urlopen(url)

    soup = BeautifulSoup(html,'lxml')

    div =soup.findAll('div',{'class':'inner_box'})

    df = BeautifulSoup(str(div[0]),features='lxml')

    a = df.findAll('a')

    for each in a:

        if each.string=='下载数据':

            new_html = adrees + each.get("href")

            html1 = urlopen(new_html)

            soup1 = BeautifulSoup(html1,features='lxml')

            txt = soup1.text.replace('(万元)','').replace('--','0')

            csv = open(r'F:\stock\python_data\财务分析\数据采集\%s%s.csv'%(stock_code,i),'w',encoding='utf-8') .write(txt)



            data = pd.read_csv(r'F:\stock\python_data\财务分析\数据采集\%s%s.csv'%(stock_code,i))

            list1 = list(range(int(data.columns[-2][:4]), int(data.columns[1][:4]) + 1))[::-1]

            new_list = ['报告日期']+[str(i) + '年' for i in list1] + [data.columns[-1][:4]]

            data.columns = new_list

            writer = pd.ExcelWriter(r'F:\stock\python_data\财务分析\数据采集\%s%s.xlsx' %(stock_code,i))

            data.to_excel(writer,index=False)

            writer.save()

            writer.close()



            p_list = os.listdir(r'F:\stock\python_data\财务分析\数据采集')

            for name in p_list:

                if name.endswith('.csv'):

                    os.remove(r'F:\stock\python_data\财务分析\数据采集\%s%s.csv'%(stock_code,i))





                #读取已有数据并按照所需的列数进行合并

path_list = os.listdir(r'F:\stock\python_data\财务分析\数据采集')

data = pd.DataFrame()

for path in path_list:

    fp = r'F:\stock\python_data\财务分析\数据采集\%s' % (path)

    # 读取excel，并指定cols为y_list的值
    dfs = pd.read_excel(fp,None,usecols=y_list)
    # 获取所有sheet
    keys = dfs.keys()

    for i in keys:
        # 获取单个sheet
        df1 = dfs[i]
        # 表格合并
        data = pd.concat([df1,data])
        # 填充缺失值为0
        data = data.fillna(0)
        # 移除指定字符
        data['报告日期'] =data['报告日期'].str.strip()
        # 返回删除了重复行的 DataFrame。
        data = data.drop_duplicates(subset=['报告日期'],keep='first')

# 使用现有列设置 DataFrame 索引。
data = data.set_index('报告日期')
# 从行或列中删除指定的标签， axis=1：从索引为1的列开始
hb = data.drop('%s年'% (int(st)-1),axis=1)





#定义所需函数

#财务指标函数定义

def hs1(cs1,cs2):

    return round(cs1/cs2,2)



def hs2(cs1,cs2,cs3):

    return round((cs1-cs2)/cs3,2)



def hs3(cs1,cs2,cs3):

    return round((cs1+cs2)/cs3,2)



def hs4(cs1,cs2,cs3):

    return round(cs1/(cs2-cs3),2)



def hs5(cs1,cs2,cs3):

    return round(cs1/(cs2+cs3),2)



def hs6(cs1,cs2,cs3,cs4):

    return round((cs1-cs2-cs3)/cs4,2)



def hs7(cs1,cs2):

    return round((cs1-cs2)/cs1,2)



def change(x,y):

    if x.loc[y,'%s年' % et] < x.loc[y,'%s年' % st]:

        return ('%s由%s年的%.2f万元，变化为%s年的%.2f万元，变化数为%.2f万元,变化率为%.2f' %

                (y,st,x.loc[y,'%s年' % st],et,x.loc[y,'%s年' % et],round(x.loc[y,'%s年' % et]-x.loc[y,'%s年' % st],2),

                 round((x.loc[y,'%s年' % et]-x.loc[y,'%s年' % st])/x.loc[y,'%s年' % st],2)))

    else:

        return ('%s由%s年的%.2f万元，变化为%s年的%.2f万元，变化数为%.2f万元,变化率为%.2f' %

                (y,st,x.loc[y,'%s年' % st],et,x.loc[y,'%s年' % et],round(x.loc[y,'%s年' % et] - x.loc[y,'%s年' % st],2),

                 round((x.loc[y, '%s年' % et] - x.loc[y, '%s年' % st])/x.loc[y, '%s年' % st],2)))



def up(x,y):

    if x.loc[y,'%s年' % et] < x.loc[y,'%s年' % st]:

        return ('%s指标由%s年的%.2f，变化为%s年的%.2f，变化数为%.2f，由此看出该公司的%s指标变弱' %

                (y,st,x.loc[y,'%s年' % st],et,x.loc[y,'%s年' % et],round(x.loc[y,'%s年' % et]-x.loc[y,'%s年' % st],2),y))

    else:

        return ('%s指标由%s年的%.2f，变化为%s年的%.2f，变化数为%.2f，由此看出该公司的%s指标变强' %

                (y,st,x.loc[y, '%s年' % st],et,x.loc[y, '%s年' % et],round(x.loc[y,'%s年' % et]-x.loc[y,'%s年'% st],2),y))



def down(x,y):

    if x.loc[y,'%s年' % et] < x.loc[y,'%s年' % st]:

        return ('%s指标由%s年的%.2f，变化为%s年的%.2f，变化数为%.2f，由此看出该公司的%s指标变强' %

                (y,st,x.loc[y,'%s年' % st],et,x.loc[y,'%s年' % et],round(x.loc[y,'%s年' % et]-x.loc[y,'%s年' % st],2),y))

    else:

        return ('%s指标由%s年的%.2f，变化为%s年的%.2f，变化数为%.2f，由此看出该公司的%s指标变弱' %

                (y,st,x.loc[y, '%s年' % st],et,x.loc[y, '%s年' % et],round(x.loc[y,'%s年' % et]-x.loc[y,'%s年'% st],2),y))

    #计算平均资产总额

def pj_data(xm):

    pj_data = []

    for i in range(len(xm)-1):

        if i < len(xm):

            pjz = round((xm[i] + xm[i+1])/2,2)

            pj_data.append(pjz)

    return pj_data



#计算增长率

def zzl_data(xm):

    zzl_data = []

    for i in range(len(xm)-1):

        if i < len(xm):

            zzl = round((xm[i]-xm[i+1])/xm[i+1],2)

            zzl_data.append(zzl)

    return zzl_data



#画图函数

def fig(x,y):

    fig = plt.figure()

    for i in range(len(x.index)):

        plt.plot(x.columns.values,x.iloc[i],label =x.index[i])

    plt.title(y)

    plt.legend(loc = 'upper right')

    fig.savefig(r'F:\stock\python_data\财务分析\成果展示\%s.png'% y)

    return fig



#生成表格函数

def data_table(x,y):

    t = d1.add_table(rows=y.shape[0] + 1, cols=y.shape[1] + 1, style='Table Grid')

    for n in range(len(x)):

        t.cell(0, n).text = x[n]

        for m in range(len(list(y.index))):

            t.cell(m + 1, 0).text = list(y.index)[m]



    for j in range(y.shape[0]):

        for i in range(y.shape[1]):

            t.cell(j + 1, i + 1).text = str(y.iloc[j, i])

    return t



#计算所需财务数据

#财报主要财务数据

zy = hb.loc[['流动资产合计','非流动资产合计','资产总计','流动负债合计','非流动负债合计','负债合计','所有者权益(或股东权益)合计','营业总收入',

             '营业总成本','利润总额','净利润']]



# #资产结构主要数据

zc = hb.loc[['货币资金','应收账款','预付款项','其他应收款','存货','流动资产合计','长期股权投资','固定资产','在建工程','无形资产',

             '长期待摊费用','非流动资产合计']]



zc2 = hb.loc[['货币资金','应收账款','预付款项','其他应收款','存货','长期股权投资','固定资产','在建工程','无形资产',

              '长期待摊费用']]



zc_bh = round(abs(zc2['%s年'%et]-zc2['%s年'%st]),2)



zc1 = round(hb.loc[['货币资金','应收账款','预付款项','其他应收款','存货','长期股权投资','固定资产','在建工程','无形资产',

                    '长期待摊费用']]/hb.loc['资产总计'],2)



#负债结构主要数据

fz = hb.loc[['短期借款','应付账款','预收账款','应付职工薪酬','应交税费','其他应付款','流动负债合计','长期借款','长期应付款',

             '非流动负债合计']]



fz2 = hb.loc[['短期借款','应付账款','预收账款','应付职工薪酬','应交税费','其他应付款','长期借款','长期应付款']]



fz_bh = round(abs(fz2['%s年'%et]-fz2['%s年'%st]),2)





fz1 = round(hb.loc[['短期借款','应付账款','预收账款','应付职工薪酬','应交税费','其他应付款','长期借款','长期应付款']]/hb.loc['负债合计'],2)



#股本结构主要数据

qy = hb.loc[['实收资本(或股本)','资本公积','盈余公积','未分配利润','所有者权益(或股东权益)合计']]



qy2 = hb.loc[['实收资本(或股本)','资本公积','盈余公积','未分配利润']]



qy_bh = round(abs(qy2['%s年'%et]-qy2['%s年'%st]),2)



qy1 = round(hb.loc[['实收资本(或股本)','资本公积','盈余公积','未分配利润']]/hb.loc['所有者权益(或股东权益)合计'],2)



#利润表主要数据

lr = hb.loc[['营业总收入','营业收入','其他业务收入','营业总成本','营业成本','其他业务成本','销售费用','管理费用','财务费用','其他业务利润',

             '营业利润','利润总额','所得税费用','净利润']]



#现金流量表主要数据

xj = hb.loc[['经营活动现金流入小计','经营活动现金流出小计','经营活动产生的现金流量净额','投资活动现金流入小计','投资活动现金流出小计',

             '投资活动产生的现金流量净额','筹资活动现金流入小计','筹资活动现金流出小计','筹资活动产生的现金流量净额','现金及现金等价物净增加额']]





#计算偿债能力指标

ldbl=hs1(cs1=hb.loc['流动资产合计'],cs2=hb.loc['流动负债合计'])                #流动比率

sdbl=hs2(cs1=hb.loc['流动资产合计'],cs2=hb.loc['存货'],cs3=hb.loc['流动负债合计'])          #速动比率

zcfzl=hs1(cs1=hb.loc['负债合计'],cs2=hb.loc['资产总计'])                                   #资产负债率

gdqybl=hs1(cs1=hb.loc['所有者权益(或股东权益)合计'],cs2=hb.loc['资产总计'])                            #股东权益比率

cqfzbl=hs1(cs1=hb.loc['非流动负债合计'],cs2=hb.loc['资产总计'])                             #长期负债比率

cqzwyu=hs4(cs1=hb.loc['非流动负债合计'],cs2=hb.loc['流动资产合计'],cs3=hb.loc['流动负债合计'])   #长期债务与营运资金比率

fzsy=hs1(cs1=hb.loc['负债合计'],cs2=hb.loc['所有者权益(或股东权益)合计'])                                   #负债与所有者权益比率

czcz=hs5(cs1=hb.loc['非流动负债合计'],cs2=hb.loc['所有者权益(或股东权益)合计'],cs3=hb.loc['非流动负债合计'])   #长期资产与长期资金比率

zbhl=hs1(cs1=hb.loc['非流动负债合计'],cs2=hb.loc['所有者权益(或股东权益)合计'])                               #资本化比率

zbgdh=hs2(cs1=hb.loc['资产总计'],cs2=hb.loc['非流动负债合计'],cs3=hb.loc['所有者权益(或股东权益)合计'])         #资本固定化比率

cqbl=hs1(cs1=hb.loc['负债合计'],cs2=hb.loc['所有者权益(或股东权益)合计'])                                     #产权比率



cz=pd.DataFrame({'流动比率':ldbl,'速动比率':sdbl,'资产负债率':zcfzl,'股东权益比率':gdqybl,'长期负债比率':cqfzbl,

                 '长期债务与营运资金比率':cqzwyu,'负债与所有者权益比率':fzsy,

                 '长期资产与长期资金比率':czcz,'资本化率':zbhl,'资本固定化比率':zbgdh,'产权比率':cqbl})

cz = cz.T



cz_bh = cz.std(axis = 1)



# 计算盈利能力指标

zzclr = hs1(cs1 = hb.loc['利润总额'],cs2 = pj_data(xm= data.loc['资产总计']))

zzcjlr = hs1(cs1 = hb.loc['净利润'],cs2 =pj_data(xm= data.loc['资产总计']))

yylr =hs1(cs1 = hb.loc['营业利润'],cs2 =hb.loc['营业总收入'] )

jzcsy = hs1(cs1 = hb.loc['净利润'],cs2 =hb.loc['所有者权益(或股东权益)合计'] )

gbbc = hs1(cs1 = hb.loc['净利润'],cs2 =hb.loc['实收资本(或股本)'] )

xsml = hs7(cs1=hb.loc['营业收入'],cs2=hb.loc['营业成本'])



yl = pd.DataFrame({'总资产利润率':zzclr,'总资产净利润率':zzcjlr,'营业利润率':yylr,'净资产收益率':jzcsy,'股本报酬率':gbbc,'销售毛利率':xsml})

yl = yl.T

yl_bh = yl.std(axis = 1)



#计算运营能力指标

yszkzzl = hs1(cs1=hb.loc['营业收入'],cs2=pj_data(xm=data.loc['应收账款']))   # 应收账款周转率

yszkzzt = round(360/yszkzzl,2)    # 应收账款周转天数

chzzl = hs1(cs1=hb.loc['营业成本'],cs2=pj_data(xm=data.loc['存货']))   # 存货周转率

chzzt = round(360/chzzl,2)    # 存货周转天数

zzzzzl = hs1(cs1=hb.loc['营业总收入'],cs2=pj_data(xm=data.loc['资产总计']))    # 总资产周转率

zzzzzt = round(360/zzzzzl,2)    # 总资产周转天数

ldzczzl = hs1(cs1=hb.loc['营业收入'],cs2=pj_data(xm=data.loc['流动资产合计']))   # 流动资产周转率

ldzczzt = round(360/ldzczzl,2)       # 流动资产周转天数



yy = pd.DataFrame({'应收账款周转率':yszkzzl,'应收账款周转天数':yszkzzt,'存货周转率':chzzl,'存货周转天数':chzzt,

                   '总资产周转率':zzzzzl,'总资产周转天数':zzzzzt,'流动资产周转率':ldzczzl,'流动资产周转天数':ldzczzt})

yy = yy.T

yy_bh = yy.std(axis = 1)



# 计算成长能力指标

zyyw = zzl_data(xm=data.loc['营业收入'])  # 主营业务收入增长率

jlrzz = zzl_data(xm=data.loc['净利润'])  # 净利润增长率

jzzzz = zzl_data(xm=data.loc['所有者权益(或股东权益)合计']) # 净资产增长率

zzzzz = zzl_data(xm=data.loc['资产总计']) # 总资产增长率



czn = pd.DataFrame({'主营业务收入增长率':zyyw,'净利润增长率':jlrzz,'净资产增长率':jzzzz,'总资产增长率':zzzzz})

czn = czn.T

czn.columns = [str(m)+'年' for m in list(range(int(st),int(et)+1))][::-1]

czn_bh = czn.std(axis = 1)



writer = pd.ExcelWriter(r'F:\stock\python_data\财务分析\成果展示\%s%s年至%s年财务分析基数数据表.xlsx' %(c_name,st,et))

zy.to_excel(writer,sheet_name='财报主要数据表')

zc.to_excel(writer,sheet_name='资产结构表')

fz.to_excel(writer,sheet_name='负债结构表')

qy.to_excel(writer,sheet_name='股本结构表')

lr.to_excel(writer,sheet_name='利润表主要数据表')

xj.to_excel(writer,sheet_name='现金流量表主要数据表')

cz.to_excel(writer,sheet_name='偿债能力主要数据表')

yl.to_excel(writer,sheet_name='盈利能力主要数据表')

yy.to_excel(writer,sheet_name='运营能力主要数据表')

czn.to_excel(writer,sheet_name='成长能力主要数据表')

writer.save()

writer.close()



#制作各种所需图

#财务主要数据变化趋势图

fig(x=zy,y='财务主要数据变化趋势图')



#资产结构变化趋势图

fig(x=zc,y='资产结构变化趋势图')



#负债结构变化趋势图

fig(x=fz,y='负债结构变化趋势图')



#股本结构变化趋势图

fig(x=qy,y='股本结构变化趋势图')



#利润表主要数据变化趋势图

fig(x=lr,y='利润表主要数据变化趋势图')



#现金流量表主要数据变化趋势图

fig(x=xj,y='现金流量表主要数据变化趋势图')



#偿债能力主要指标变化趋势图

fig(x=cz,y='偿债能力主要指标变化趋势图')



#盈利主要指标变化趋势图

fig(x=yl,y='盈利能力主要指标变化趋势图')



#运营能力主要指标变化趋势图

fig(x=yy,y='运营能力主要指标变化趋势图')



#成长能力主要指标变化趋势图

fig(x=czn,y='成长能力主要指标变化趋势图')





#文章正文部分

#文章正文标题部分

d1 = Document()

p1 =d1.add_paragraph()

p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

r1 = p1.add_run('%s财务分析报告' % c_name )

r1.font.size = Pt(14)

r1.font.bold = True

r1.font.name = '仿宋_GB2312'

r1._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')



#文章正文第一部分——财报主要数据

p2 =d1.add_paragraph()

r2 = p2.add_run('   一、公司基本财务数据 \n   下表数据为%s%s年至%s年主要财务数据：'%(c_name,st,et))

r2.font.size = Pt(10)

# r2.font.bold = True

r2.font.name = '仿宋_GB2312'

r2._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')



#插入财报主要数据表格

data_table(x=y_list,y=zy)



#财报主要数据的描述

p22 = d1.add_paragraph()

run22 = p22.add_run('    从上表数据中可以看出，%s；%s；%s；%s；%s；%s；%s。\n     下图为各项数据的变化趋势图：' % (change(x=zy,y = '资产总计'),

                                                                                   change(x=zy,y = '负债合计'),change(x=zy,y = '所有者权益(或股东权益)合计'),change(x=zy,y = '营业总收入'),change(x=zy,y = '营业总成本'),

                                                                                   change(x=zy,y = '利润总额'),change(x=zy,y = '净利润')))

run22.font.name = "仿宋_GB2312"

run22._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')

run22.font.size = Pt(10)



#插入财报主要数据变化趋势图

tu1 = d1.add_picture(r'F:\stock\python_data\财务分析\成果展示\财务主要数据变化趋势图.png',width=Inches(6),height=Inches(4))

tu1.alignment = WD_ALIGN_PARAGRAPH.CENTER

#

#文章正文第二部分——资产结构

p3 =d1.add_paragraph()

r3 = p3.add_run('   二、资产负债表分析 \n   1、资产结构分析\n    下表数据为%s%s年至%s年资产结构分析主要数据。'%(c_name,st,et))

r3.font.size = Pt(10)

# r3.font.bold = True

r3.font.name = '仿宋_GB2312'

r3._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')



#插入资产结构表格

data_table(x=y_list,y=zc)



p31 = d1.add_paragraph()

run31 = p31.add_run('    '

                    '从上表数据中可以看出，在%s年至%s年度中%s；%s；其中变化最大的项目为%s；变化最小的项目为%s。'

                    '在%s年度的资产结构中占比最大的项目%s，占比比例为%.2f，此项目%s年度的占比为%.2f；'

                    '占比最小的项目为%s其中占比为%.2f，此项目%s年度的占比为%.2f。'

                    '以上是对资产结构的整体介绍，其他具体项目的变化趋势见下图：' %(

                        st,et,

                        change(x=zc,y = '流动资产合计'),change(x=zc,y = '非流动资产合计'),change(x=zc2,y = zc_bh.idxmax()),change(x=zc2,y = zc_bh.idxmin()),

                        et,zc1['%s年'%et].idxmax(),zc1.loc[zc1['%s年'%et].idxmax(),'%s年'%et],st,zc1.loc[zc1['%s年'%st].idxmax(),'%s年'%st],

                        zc1['%s年'%et].idxmin(),zc1.loc[zc1['%s年'%et].idxmin(),'%s年'%et],st,zc1.loc[zc1['%s年'%st].idxmin(),'%s年'%st]))

run31.font.name = "仿宋_GB2312"

run31._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')

run31.font.size = Pt(10)



#

tu2 = d1.add_picture(r'F:\stock\python_data\财务分析\成果展示\资产结构变化趋势图.png',width=Inches(6),height=Inches(4))

tu2.alignment = WD_ALIGN_PARAGRAPH.CENTER





p4 =d1.add_paragraph()

r4 = p4.add_run('   2、负债结构分析\n    下表数据为%s%s年至%s年负债结构分析主要数据。'%(c_name,st,et))

r4.font.size = Pt(10)

# r4.font.bold = True

r4.font.name = '仿宋_GB2312'

r4._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')

#

#插入负债结构表

data_table(x=y_list,y=fz)



p41 = d1.add_paragraph()

run41 = p41.add_run('    '

                    '从上表数据中可以看出，在%s年至%s年度中%s；%s；其中变化最大的项目为%s；变化最小的项目为%s。'

                    '在%s年度的负债结构中占比最大的项目%s，占比比例为%.2f，此项目%s年度的占比为%.2f；'

                    '占比最小的项目为%s其中占比为%.2f，此项目%s年度的占比为%.2f。'

                    '以上是对负债结构的整体介绍，其他具体项目的变化趋势见下图：' %(

                        st,et,

                        change(x=fz,y='流动负债合计'),change(x=fz,y='非流动负债合计'),change(x=fz2,y=fz_bh.idxmax()),change(x=fz2,y=fz_bh.idxmin()),

                        et,fz1['%s年'%et].idxmax(),fz1.loc[fz1['%s年'%et].idxmax(),'%s年'%et],st,fz1.loc[fz1['%s年'%st].idxmax(),'%s年'%st],

                        fz1['%s年'%et].idxmax(),fz1.loc[fz1['%s年'%et].idxmax(),'%s年'%et],st,fz1.loc[fz1['%s年'%st].idxmax(),'%s年'%st]

                    ))

run41.font.name = "仿宋_GB2312"

run41._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')

run41.font.size = Pt(10)



tu3 = d1.add_picture(r'F:\stock\python_data\财务分析\成果展示\负债结构变化趋势图.png',width=Inches(6),height=Inches(4))

tu3.alignment = WD_ALIGN_PARAGRAPH.CENTER





p5 =d1.add_paragraph()

r5 = p5.add_run('   3、所有者权益结构分析\n    下表数据为%s%s年至%s年所有者权益结构分析主要数据。'%(c_name,st,et))

r5.font.size = Pt(10)

# r5.font.bold = True

r5.font.name = '仿宋_GB2312'

r5._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')



#插入权益结构表

data_table(x=y_list,y=qy)



p51 = d1.add_paragraph()

run51 = p51.add_run('    '

                    '从上表数据中可以看出，在%s年至%s年度中%s；%s；其中变化最大的项目为%s；变化最小的项目为%s。'

                    '在%s年度的股本结构中占比最大的项目%s，占比比例为%.2f，此项目%s年度的占比为%.2f；'

                    '占比最小的项目为%s其中占比为%.2f，此项目%s年度的占比为%.2f。'

                    '以上是对股本结构的整体介绍，其他具体项目的变化趋势见下图：' %(

                        st,et,

                        change(x=qy, y='实收资本(或股本)'), change(x=qy, y='未分配利润'), change(x=qy2, y=qy_bh.idxmax()),change(x=qy2, y=qy_bh.idxmin()),

                        et,qy1['%s年'%et].idxmax(),qy1.loc[qy1['%s年'%et].idxmax(),'%s年'%et],st,qy1.loc[qy1['%s年'%st].idxmax(),'%s年'%st],

                        qy1['%s年'%et].idxmax(),qy1.loc[qy1['%s年'%et].idxmax(),'%s年'%et],st,qy1.loc[qy1['%s年'%st].idxmax(),'%s年'%st]

                    ))

run51.font.name = "仿宋_GB2312"

run51._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')

run51.font.size = Pt(10)



tu4 = d1.add_picture(r'F:\stock\python_data\财务分析\成果展示\股本结构变化趋势图.png',width=Inches(6),height=Inches(4))

tu4.alignment = WD_ALIGN_PARAGRAPH.CENTER



#

p6 =d1.add_paragraph()

r6 = p6.add_run('   三、利润表分析\n    下表数据为%s%s年至%s年利润表主要数据。'%(c_name,st,et))

r6.font.size = Pt(10)

# r6.font.bold = True

r6.font.name = '仿宋_GB2312'

r6._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')



#插入利润表

data_table(x=y_list,y=lr)



p61 = d1.add_paragraph()

run61 = p61.add_run('    从上表数据中可以看出，%s;%s;%s;%s;%s;%s;%s;%s;%s;%s;%s;%s;%s;%s。\n  下图为各项数据的变化趋势图：'%

                    (change(x=lr,y='营业总收入'),change(x=lr,y='营业收入'),change(x=lr,y='其他业务收入'),change(x=lr,y='营业总成本'),

                     change(x=lr,y='营业成本'),change(x=lr,y='其他业务成本'),change(x=lr,y='销售费用'),change(x=lr,y='管理费用'),

                     change(x=lr,y='财务费用'),change(x=lr,y='其他业务利润'),change(x=lr,y='营业利润'),change(x=lr,y='利润总额'),

                     change(x=lr,y='所得税费用'),change(x=lr,y='净利润'),))



run61.font.name = "仿宋_GB2312"

run61._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')

run61.font.size = Pt(10)



tu5 = d1.add_picture(r'F:\stock\python_data\财务分析\成果展示\利润表主要数据变化趋势图.png',width=Inches(6),height=Inches(4))

tu5.alignment = WD_ALIGN_PARAGRAPH.CENTER

#



p7 =d1.add_paragraph()

r7 = p7.add_run('   四、现金流量表分析\n    下表数据为%s%s年至%s年现金流量表主要数据。'%(c_name,st,et))

r7.font.size = Pt(10)

# 7.font.bold = True

r7.font.name = '仿宋_GB2312'

r7._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')



#插入现金流量表

data_table(x=y_list,y=xj)



p71 = d1.add_paragraph()

run71 = p71.add_run('    从上表数据中可以看出，%s;%s;%s;%s;%s;%s;%s;%s;%s;%s。\n  下图为各项数据的变化趋势图：' % (

    change(x=xj,y='经营活动现金流入小计'),change(x=xj,y='经营活动现金流出小计'),change(x=xj,y='经营活动产生的现金流量净额'),

    change(x=xj,y='投资活动现金流入小计'),change(x=xj,y='投资活动现金流出小计'),change(x=xj,y='投资活动产生的现金流量净额'),

    change(x=xj,y='筹资活动现金流入小计'),change(x=xj,y='筹资活动现金流出小计'),change(x=xj,y='筹资活动产生的现金流量净额'),

    change(x=xj,y='现金及现金等价物净增加额')))



run71.font.name = "仿宋_GB2312"

run71._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')

run71.font.size = Pt(10)

#

tu6 = d1.add_picture(r'F:\stock\python_data\财务分析\成果展示\现金流量表主要数据变化趋势图.png',width=Inches(6),height=Inches(4))

tu6.alignment = WD_ALIGN_PARAGRAPH.CENTER

#

#

p8 =d1.add_paragraph()

r8 = p8.add_run('   五、综合财务指标分析 \n   1、偿债能力分析\n    下表数据为%s%s年至%s年尝债能力分析指标主要数据。'%(c_name,st,et))

r8.font.size = Pt(10)

# r8.font.bold = True

r8.font.name = '仿宋_GB2312'

r8._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')



#插入偿债能力表

data_table(x=y_list,y=cz)



p81 = d1.add_paragraph()

run81 = p81.add_run('    上表数据反映出，%s；%s；%s；%s。其中变化最大的指标为%s。下图为各项指标数据的变化趋势图：'%(

    up(x=cz,y='流动比率'),up(x=cz,y='速动比率'),down(x=cz,y='资产负债率'),down(x=cz,y='产权比率'),change(x=cz, y=cz_bh.idxmax())

) )

run81.font.name = "仿宋_GB2312"

run81._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')

run81.font.size = Pt(10)



tu7 = d1.add_picture(r'F:\stock\python_data\财务分析\成果展示\偿债能力主要指标变化趋势图.png',width=Inches(6),height=Inches(4))

tu7.alignment = WD_ALIGN_PARAGRAPH.CENTER



p9 =d1.add_paragraph()

r9 = p9.add_run('   2、盈利能力分析\n    下表数据为%s%s年至%s年盈利能力分析指标主要数据。'%(c_name,st,et))

r9.font.size = Pt(10)

# r9.font.bold = True

r9.font.name = '仿宋_GB2312'

r9._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')



#插入盈利能力表

data_table(x=y_list,y=yl)



p91 = d1.add_paragraph()

run91 = p91.add_run('     上表数据反映出，%s；%s；%s；%s；%s；%s。其中变化最大的指标为%s。下图为各项指标数据的变化趋势图：'%(

    up(x=yl,y='总资产利润率'),up(x=yl,y='总资产净利润率'),up(x=yl,y='营业利润率'),up(x=yl,y='净资产收益率'),up(x=yl,y='股本报酬率'),

    up(x=yl,y='销售毛利率'),change(x=yl, y=yl_bh.idxmax())

) )

run91.font.name = "仿宋_GB2312"

run91._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')

run91.font.size = Pt(10)



tu8 = d1.add_picture(r'F:\stock\python_data\财务分析\成果展示\盈利能力主要指标变化趋势图.png',width=Inches(6),height=Inches(4))

tu8.alignment = WD_ALIGN_PARAGRAPH.CENTER



p10 =d1.add_paragraph()

r10 = p10.add_run('   3、运营能力分析\n    下表数据为%s%s年至%s年运营能力分析指标主要数据。'%(c_name,st,et))

r10.font.size = Pt(10)

# r10.font.bold = True

r10.font.name = '仿宋_GB2312'

r10._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')



#插入运营能力表

data_table(x=y_list,y=yy)



p101 = d1.add_paragraph()

run101 = p101.add_run('    上表数据反映出，%s；%s；%s；%s；%s；%s；%s；%s。下图为各项指标数据的变化趋势图：' %(

    up(x=yy,y='应收账款周转率'),down(x=yy,y='应收账款周转天数'),up(x=yy,y='存货周转率'),down(x=yy,y='存货周转天数'),

    up(x=yy,y='总资产周转率'),down(x=yy,y='总资产周转天数'),up(x=yy,y='流动资产周转率'),down(x=yy,y='流动资产周转天数')))

run101.font.name = "仿宋_GB2312"

run101._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')

run101.font.size = Pt(10)



tu9 = d1.add_picture(r'F:\stock\python_data\财务分析\成果展示\运营能力主要指标变化趋势图.png',width=Inches(6),height=Inches(4))

tu9.alignment = WD_ALIGN_PARAGRAPH.CENTER



p11 =d1.add_paragraph()

r11 = p11.add_run('   4、成长能力分析\n    下表数据为%s%s年至%s年成长能力分析指标主要数据。'%(c_name,st,et))

r11.font.size = Pt(10)

# r11.font.bold = True

r11.font.name = '仿宋_GB2312'

r11._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')



#插入成长能力表

data_table(x=y_list,y=czn)



p111 = d1.add_paragraph()

run111 = p111.add_run('    上表数据反映出，%s;%s;%s;%s。\n    下图为各项指标数据的变化趋势图：' %(

    up(x=czn,y='主营业务收入增长率'),up(x=czn,y='净利润增长率'),up(x=czn,y='净资产增长率'),up(x=czn,y='总资产增长率')))

run111.font.name = "仿宋_GB2312"

run111._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')

run111.font.size = Pt(10)



tu10 = d1.add_picture(r'F:\stock\python_data\财务分析\成果展示\成长能力主要指标变化趋势图.png',width=Inches(6),height=Inches(4))

tu10.alignment = WD_ALIGN_PARAGRAPH.CENTER





d1.save(r'F:\stock\python_data\财务分析\成果展示\%s%s年度至%s年度财务分析.docx' % (c_name,st,et))

